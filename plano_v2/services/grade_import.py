"""Leitura conservadora de grades curriculares, sempre sem persistência.

Os formatos de arquivo e a colagem de texto passam pela mesma normalização para
que o usuário revise horas, status e datas antes de confirmar uma importação.
"""
from __future__ import annotations

import csv
import io
import re
import unicodedata
import zipfile
from datetime import date, datetime
from pathlib import Path


ALLOWED = {"xlsx", "csv", "txt", "docx", "pdf"}
MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_XLSX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_ROWS = 1000
MAX_COLUMNS = 30
HEADER_SCAN_ROWS = 20
PREFERRED_SHEET = "GRADE_PARA_IMPORTAR"
IGNORED_SHEETS = {"INSTRUCOES", "OPCOES ELETIVAS"}
ACADEMIC_STATUSES = {"not_available", "available", "in_progress", "completed", "failed", "locked", "exempted"}


def _clean(value):
    return re.sub(r"\s+", " ", str(value or "").replace("\xa0", " ")).strip(" -;|")


def normalized(value):
    text = unicodedata.normalize("NFKD", _clean(value)).casefold()
    text = "".join(character for character in text if not unicodedata.combining(character))
    return re.sub(r"\s+", " ", text).strip()


def normalized_name(value):
    """Chave estável de nome para alertar duplicatas sem depender do SQLite."""
    return re.sub(r"[^a-z0-9]+", " ", normalized(value)).strip()


def _header_key(value):
    return re.sub(r"[^a-z0-9]+", " ", normalized(value)).strip()


HEADER_ALIASES = {
    "name": {"disciplina", "materia", "nome", "unidade curricular", "componente curricular"},
    "code": {"codigo", "sigla", "cod"},
    "period": {"periodo", "modulo", "modulo periodo", "semestre", "etapa", "ciclo", "ucfc"},
    "workload_hours": {"carga horaria", "carga horaria h", "ch", "horas", "ch total"},
    "academic_status": {"status", "situacao"},
    "sort_order": {"ordem", "sequencia", "posicao"},
    "start_date": {"data de inicio", "inicio"},
    "end_date": {"data de termino", "termino", "fim"},
    "notes": {"observacoes", "notas", "comentarios"},
    "include": {"importar", "incluir", "selecionar"},
}
HEADER_LOOKUP = {alias: field for field, aliases in HEADER_ALIASES.items() for alias in aliases}
STATUS_ALIASES = {
    "nao iniciado": "not_available", "not available": "not_available", "not_available": "not_available",
    "disponivel": "available", "available": "available",
    "em andamento": "in_progress", "in progress": "in_progress", "in_progress": "in_progress",
    "concluida": "completed", "concluido": "completed", "completed": "completed",
    "reprovada": "failed", "reprovado": "failed", "failed": "failed",
    "bloqueada": "locked", "bloqueado": "locked", "locked": "locked",
    "dispensada": "exempted", "dispensado": "exempted", "exempted": "exempted",
}
ADMINISTRATIVE = re.compile(r"\b(conselho de classe|rematricula|recesso escolar|periodo de exame|exame final|ch total)\b", re.I)


def _headers(headers):
    mapping, duplicates = {}, []
    for index, header in enumerate(headers):
        field = HEADER_LOOKUP.get(_header_key(header))
        if not field:
            continue
        if field in mapping:
            duplicates.append(_clean(header))
            continue
        mapping[field] = index
    return mapping, duplicates


def _as_date(value, label):
    if value in (None, ""):
        return None, None
    if isinstance(value, datetime):
        return value.date().isoformat(), None
    if isinstance(value, date):
        return value.isoformat(), None
    text = _clean(value)
    try:
        return date.fromisoformat(text).isoformat(), None
    except ValueError:
        pass
    for pattern in ("%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(text, pattern).date().isoformat(), f"{label} foi convertido de {text} para uma data ISO."
        except ValueError:
            continue
    return None, f"{label} deve usar uma data válida (AAAA-MM-DD)."


def _hours(value):
    if value in (None, ""):
        return None, None
    text = _clean(value).lower().replace(",", ".")
    match = re.fullmatch(r"(\d+(?:\.\d+)?)\s*(?:h|hs|hora|horas)?", text)
    if not match or float(match.group(1)) <= 0:
        return None, "Carga horária deve conter apenas um número positivo de horas."
    hours = float(match.group(1))
    return int(hours) if hours.is_integer() else hours, None


def _minutes(value):
    if value in (None, ""):
        return None, None
    try:
        minutes = int(value)
    except (TypeError, ValueError):
        return None, "Carga em minutos deve ser um número inteiro positivo."
    return (minutes, None) if minutes > 0 else (None, "Carga em minutos deve ser maior que zero.")


def _include(value):
    if value in (None, ""):
        return True, None
    answer = normalized(value)
    if answer in {"sim", "s", "true", "1", "yes"}:
        return True, None
    if answer in {"nao", "n", "false", "0", "no"}:
        return False, None
    return False, "Valor de Importar? não reconhecido; a linha ficou desmarcada por segurança."


def _status(value):
    if value in (None, ""):
        return "not_available", None, None
    key = normalized(value)
    mapped = STATUS_ALIASES.get(key)
    if mapped:
        return mapped, None, None
    return None, f"Status “{_clean(value)}” não reconhecido.", "Escolha um status válido antes de importar esta linha."


def _order(value, fallback):
    if value in (None, ""):
        return fallback, None
    try:
        order = int(value)
    except (TypeError, ValueError):
        return fallback, "Ordem deve ser um número inteiro igual ou maior que zero."
    return (order, None) if order >= 0 else (fallback, "Ordem deve ser igual ou maior que zero.")


def _append_note(current, extra):
    values = [_clean(current), _clean(extra)]
    return "\n".join(value for value in values if value) or None


def _state(item):
    if not item["include"]:
        return "excluded"
    if item["errors"]:
        return "blocked"
    return "review" if item["warnings"] else "ready"


def normalize_row(raw, source, source_index, default_order, warnings=None, confidence="high"):
    """Transforma uma linha externa em um item seguro para a prévia editável."""
    row_warnings, errors = list(warnings or []), []
    include, include_warning = _include(raw.get("include"))
    if include_warning:
        row_warnings.append(include_warning)
    name = _clean(raw.get("name"))
    if not name:
        errors.append("Disciplina é obrigatória.")

    hours, hours_error = _hours(raw.get("workload_hours"))
    supplied_minutes, minutes_error = _minutes(raw.get("workload_minutes"))
    if hours_error:
        errors.append(hours_error)
    if minutes_error:
        errors.append(minutes_error)
    workload_minutes = round(hours * 60) if hours is not None else supplied_minutes
    if hours is not None and supplied_minutes is not None and workload_minutes != supplied_minutes:
        errors.append("A carga em horas não corresponde à carga em minutos.")
    if hours is None and supplied_minutes is not None:
        hours = supplied_minutes / 60
    if workload_minutes is None:
        row_warnings.append("Carga horária ausente.")

    raw_status = raw.get("academic_status", raw.get("status"))
    academic_status, status_warning, status_error = _status(raw_status)
    if status_warning:
        row_warnings.append(status_warning)
    if status_error:
        errors.append(status_error)
    sort_order, order_error = _order(raw.get("sort_order"), default_order)
    if order_error:
        errors.append(order_error)
    start_date, start_issue = _as_date(raw.get("start_date"), "Data de início")
    end_date, end_issue = _as_date(raw.get("end_date"), "Data de término")
    for issue in (start_issue, end_issue):
        if issue:
            if "deve usar" in issue:
                errors.append(issue)
            else:
                row_warnings.append(issue)
    if start_date and end_date and start_date > end_date:
        errors.append("A data de término não pode ser anterior à data de início.")

    item = {
        "include": include,
        "name": name,
        "code": _clean(raw.get("code")) or None,
        "period": _clean(raw.get("period")) or None,
        "workload_hours": hours,
        "workload_minutes": workload_minutes,
        "academic_status": academic_status,
        # Mantém o valor original para a tela conseguir distinguir um status
        # desconhecido de uma escolha manual feita durante a revisão.
        "status_raw": _clean(raw_status) or None,
        "requires_review": bool(status_error),
        "sort_order": sort_order,
        "start_date": start_date,
        "end_date": end_date,
        "notes": _clean(raw.get("notes")) or None,
        "warnings": row_warnings,
        "errors": errors,
        "source": source,
        "source_index": source_index,
        "confidence": confidence,
        "duplicate_action": "skip",
    }
    item["state"] = _state(item)
    return item


def _nonempty(values):
    return any(value not in (None, "") for value in values)


def _is_repeated_header(values):
    mapping, _ = _headers(values)
    return len(mapping) >= 2 or (mapping.get("name") == 0 and len(values) == 1)


def normalize_table(headers, rows, source_label, format_name, only_names=False):
    mapping, repeated_headers = _headers(headers)
    if "name" not in mapping:
        if only_names:
            mapping = {"name": 0}
            headers = ["Disciplina"]
        else:
            raise ValueError("Não foi encontrada uma coluna Disciplina. Use os cabeçalhos do modelo oficial.")
    known_indexes = set(mapping.values())
    unknown_headers = [(index, _clean(header)) for index, header in enumerate(headers) if index not in known_indexes and _clean(header)]
    global_warnings = []
    if repeated_headers:
        global_warnings.append("Cabeçalhos repetidos foram ignorados: " + ", ".join(repeated_headers) + ".")
    if only_names:
        global_warnings.append("Foram encontrados somente nomes; revise carga, período, código e status antes de importar.")

    items = []
    for offset, row_data in enumerate(rows, 1):
        values = list(row_data.get("values", []))
        if not _nonempty(values):
            continue
        if _is_repeated_header(values):
            continue
        raw = {field: values[index] if index < len(values) else None for field, index in mapping.items()}
        extras = []
        for index, header in unknown_headers:
            value = values[index] if index < len(values) else None
            if value not in (None, ""):
                extras.append(f"{header}: {_clean(value)}")
        if extras:
            raw["notes"] = _append_note(raw.get("notes"), "Dados adicionais — " + "; ".join(extras))
        item = normalize_row(
            raw,
            source=row_data.get("source") or source_label,
            source_index=row_data.get("source_index", offset),
            default_order=len(items) + 1,
            warnings=row_data.get("warnings"),
            confidence=row_data.get("confidence", "high"),
        )
        items.append(item)
        if len(items) > MAX_ROWS:
            raise ValueError(f"A importação aceita no máximo {MAX_ROWS} linhas.")
    if not items:
        raise ValueError("Nenhuma disciplina foi identificada. Corrija o arquivo ou cadastre manualmente.")
    return {"format": format_name, "items": items, "summary": summarize(items), "warnings": global_warnings}


def summarize(items):
    selected = [item for item in items if item.get("include")]
    valid = [item for item in selected if not item.get("errors")]
    # O total representa as linhas selecionadas na prévia, mesmo que alguma
    # precise de correção antes do botão final ser liberado.
    total_minutes = sum(int(item["workload_minutes"] or 0) for item in selected)
    return {
        "recognized": len(items),
        "selected": len(selected),
        "valid": len(valid),
        "with_warnings": sum(bool(item.get("warnings")) for item in items),
        "blocked": sum(bool(item.get("errors")) and item.get("include") for item in items),
        "duplicates": sum(bool(item.get("duplicate")) for item in items),
        "total_minutes": total_minutes,
        "total_hours": total_minutes / 60,
    }


def _decode(content):
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _text_rows(text):
    lines = [line for line in text.splitlines() if line.strip()]
    if not lines:
        return [], [], True
    sample = "\n".join(lines[:8])
    delimiter = "\t" if "\t" in sample else None
    if delimiter is None:
        candidates = [candidate for candidate in (";", ",") if candidate in sample]
        if candidates:
            try:
                delimiter = csv.Sniffer().sniff(sample, delimiters=";,\t").delimiter
            except csv.Error:
                delimiter = candidates[0]
    if delimiter:
        values = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    elif any(re.search(r"\s{2,}", line) for line in lines[:8]):
        values = [re.split(r"\s{2,}", line.strip()) for line in lines]
    else:
        values = [[line.strip()] for line in lines]
    headers = values[0]
    mapping, _ = _headers(headers)
    if "name" in mapping:
        return headers, [{"values": row, "source_index": index + 2} for index, row in enumerate(values[1:])], False
    return ["Disciplina"], [{"values": [row[0] if row else ""], "source_index": index + 1} for index, row in enumerate(values)], True


def preview_text(text, source_label="Texto colado", format_name="paste"):
    if len(text) > MAX_UPLOAD_BYTES:
        raise ValueError("O texto colado excede o limite de 5 MB.")
    headers, rows, only_names = _text_rows(text)
    return normalize_table(headers, rows, source_label, format_name, only_names=only_names)


def _safe_xlsx(content):
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if any(entry.filename.casefold().endswith("vbaproject.bin") for entry in entries):
                raise ValueError("Planilhas com macros não são aceitas. Use um arquivo .xlsx sem macros.")
            if sum(entry.file_size for entry in entries) > MAX_XLSX_UNCOMPRESSED_BYTES:
                raise ValueError("A planilha excede o limite seguro de conteúdo descompactado.")
    except zipfile.BadZipFile as error:
        raise ValueError("O arquivo .xlsx não é uma planilha válida.") from error


def _sheet_headers(worksheet):
    for index, row in enumerate(worksheet.iter_rows(values_only=True), 1):
        if index > HEADER_SCAN_ROWS:
            break
        values = list(row)[:MAX_COLUMNS]
        mapping, _ = _headers(values)
        if "name" in mapping:
            return index, values, mapping
    return None, [], {}


def _xlsx_code(value, number_format):
    if isinstance(value, (int, float)) and float(value).is_integer() and re.fullmatch(r"0+", str(number_format or "")):
        return str(int(value)).zfill(len(str(number_format)))
    return value


def _preview_xlsx(content, sheet_name=None):
    _safe_xlsx(content)
    try:
        from openpyxl import load_workbook
        values_book = load_workbook(io.BytesIO(content), read_only=True, data_only=True, keep_links=False)
        formulas_book = load_workbook(io.BytesIO(content), read_only=True, data_only=False, keep_links=False)
    except Exception as error:
        raise ValueError("Não foi possível abrir a planilha .xlsx.") from error
    sheets = []
    for worksheet in values_book.worksheets:
        header_row, _, mapping = _sheet_headers(worksheet)
        ignored = normalized_name(worksheet.title).upper() in IGNORED_SHEETS
        sheets.append({"name": worksheet.title, "compatible": bool(mapping), "ignored": ignored, "header_row": header_row})
    compatible = [sheet for sheet in sheets if sheet["compatible"] and not sheet["ignored"]]
    preferred = next((sheet for sheet in compatible if normalized_name(sheet["name"]) == normalized_name(PREFERRED_SHEET)), None)
    if sheet_name:
        selected = next((sheet for sheet in sheets if sheet["name"] == sheet_name), None)
        if not selected or not selected["compatible"]:
            raise ValueError("A aba selecionada não contém uma coluna Disciplina compatível.")
    elif preferred:
        selected = preferred
    elif len(compatible) == 1:
        selected = compatible[0]
    elif len(compatible) > 1:
        return {
            "format": "xlsx", "items": [], "summary": summarize([]),
            "warnings": ["Há mais de uma aba compatível. Escolha qual grade deseja revisar."],
            "sheets": sheets, "selected_sheet": None, "requires_sheet_selection": True,
        }
    else:
        raise ValueError("Nenhuma aba compatível foi encontrada. Use os cabeçalhos do modelo oficial.")

    values_sheet = values_book[selected["name"]]
    formulas_sheet = formulas_book[selected["name"]]
    header_row, headers, mapping = _sheet_headers(values_sheet)
    rows = []
    for source_index, (value_cells, formula_cells) in enumerate(zip(values_sheet.iter_rows(), formulas_sheet.iter_rows()), 1):
        if source_index <= header_row:
            continue
        if source_index > header_row + MAX_ROWS:
            raise ValueError(f"A importação aceita no máximo {MAX_ROWS} linhas.")
        values = [cell.value for cell in value_cells][:MAX_COLUMNS]
        formulas = list(formula_cells)[:MAX_COLUMNS]
        warnings = []
        for field, column in mapping.items():
            if column >= len(values):
                continue
            formula = formulas[column] if column < len(formulas) else None
            if formula and formula.data_type == "f" and values[column] in (None, ""):
                warnings.append(f"Fórmula sem valor calculado em {formula.coordinate}; informe o valor manualmente.")
            if field == "code" and formula:
                values[column] = _xlsx_code(values[column], formula.number_format)
        rows.append({"values": values, "source": f"Planilha {selected['name']}", "source_index": source_index, "warnings": warnings})
    result = normalize_table(headers, rows, f"Planilha {selected['name']}", "xlsx")
    result.update({"sheets": sheets, "selected_sheet": selected["name"], "requires_sheet_selection": False})
    return result


def _pdf_noise(line):
    value = normalized(line)
    return not value or value.startswith("http") or "disciplina carga horaria" in value or bool(re.match(r"\d{2}/\d{2}/\d{4},", value)) or "pagina" in value


def _uninter_pdf(text):
    period, raw_rows = None, []
    for source_index, source_line in enumerate(text.splitlines(), 1):
        line = _clean(source_line)
        key = normalized(line)
        if _pdf_noise(line):
            continue
        if key.startswith("ucfc "):
            period = line
            continue
        if "opcoes" in key and "eletiva" in key:
            period = "OPÇÕES ELETIVAS"
            continue
        if period == "OPÇÕES ELETIVAS":
            continue
        match = re.fullmatch(r"(.+?)\s+(\d{1,4})", line)
        if match and not ADMINISTRATIVE.search(normalized(match.group(1))):
            raw_rows.append({"values": [match.group(1), period, match.group(2)], "source_index": source_index, "confidence": "medium", "warnings": ["Extraído de PDF; confira antes de importar."]})
            continue
        if "trabalho de conclusao" in key:
            raw_rows.append({"values": [line, period, None, None, None, None, None, None, None, "Não"], "source_index": source_index, "confidence": "low", "warnings": ["Carga horária ausente; a linha foi desmarcada."]})
    headers = ["Disciplina", "Módulo/Período", "Carga horária (h)", "Código", "Status", "Ordem", "Data de início", "Data de término", "Observações", "Importar?"]
    return normalize_table(headers, raw_rows, "PDF", "pdf")


def _schedule_pdf(text):
    period, pending, raw_rows = None, None, []
    total_match = re.search(r"eletrot[eéê]cnica\s*\n\s*(\d{3,4})", text, re.I)
    declared_total = int(total_match.group(1)) if total_match else None
    numeric_pattern = re.compile(r"\b(\d{1,3})\s+\d{1,3}\s+\d{1,3}\s+\d{1,3}\s+(\d{2}/\d{2}/\d{4})")
    for source_index, source_line in enumerate(text.splitlines(), 1):
        line = _clean(source_line)
        key = normalized(line)
        if not line or _pdf_noise(line):
            continue
        if key in {"ambientacao", "introdutorio", "especifico i", "especifico ii", "especifico iii"}:
            period, pending = line, None
            continue
        if ADMINISTRATIVE.search(key) and not key.startswith("recesso escolar"):
            pending = None
            continue
        match = numeric_pattern.search(line)
        if not match:
            if 3 <= len(line) <= 160 and not re.search(r"\d{2}/\d{2}/\d{4}", line) and "curso tecnico" not in key and key not in {"curso", "ch total", "modulo unidade curricular ch", "eletrotecnica"}:
                pending = line
            continue
        prefix = line[:match.start()]
        prefix = re.sub(r"^.*?\b(?:seg|ter|qua|qui|sex|sab|sáb)\.\s*", "", prefix, flags=re.I)
        prefix = re.sub(r"\([^)]*\)", "", prefix).strip()
        if prefix.startswith("Recesso escolar"):
            prefix = ""
        name = _clean(prefix) or pending
        if not name or ADMINISTRATIVE.search(normalized(name)):
            pending = None
            continue
        # O padrão já consumiu a data de início. Só depois dela vêm o término
        # e, quando existir, datas administrativas (AVA/exame). Linhas de PDF
        # podem trazer no começo o término da atividade anterior, por isso não
        # procuramos datas no prefixo.
        tail_dates = re.findall(r"\d{2}/\d{2}/\d{4}", line[match.end():])
        raw_rows.append({"values": [name, period, match.group(1), None, None, None, match.group(2), tail_dates[0] if tail_dates else None], "source_index": source_index, "confidence": "low", "warnings": ["Extraído de cronograma PDF; confira nome, datas e carga antes de importar."]})
        pending = None
    headers = ["Disciplina", "Módulo/Período", "Carga horária (h)", "Código", "Status", "Ordem", "Data de início", "Data de término"]
    result = normalize_table(headers, raw_rows, "PDF", "pdf")
    if declared_total:
        result["declared_total_hours"] = declared_total
        difference = result["summary"]["total_hours"] - declared_total
        if difference:
            result["warnings"].append(f"A carga identificada ({result['summary']['total_hours']:g} h) difere do total declarado no documento ({declared_total} h) em {difference:g} h.")
    return result


def _preview_pdf(content):
    from pypdf import PdfReader
    text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages)
    return _uninter_pdf(text) if "UCFC" in text.upper() else _schedule_pdf(text)


def _preview_docx(content):
    from docx import Document
    document = Document(io.BytesIO(content))
    table_text = ["\t".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
    text = "\n".join([paragraph.text for paragraph in document.paragraphs] + table_text)
    return preview_text(text, source_label="DOCX", format_name="docx")


def preview(file, filename, sheet_name=None):
    extension = Path(filename or "").suffix.lower().lstrip(".")
    if extension not in ALLOWED:
        raise ValueError("Formato não suportado. Use XLSX, PDF, DOCX, CSV ou TXT.")
    content = file.read(MAX_UPLOAD_BYTES + 1)
    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError("O arquivo excede o limite de 5 MB.")
    if extension == "xlsx":
        return _preview_xlsx(content, sheet_name)
    if extension == "csv":
        return preview_text(_decode(content), source_label="CSV", format_name="csv")
    if extension == "txt":
        return preview_text(_decode(content), source_label="TXT", format_name="txt")
    if extension == "docx":
        return _preview_docx(content)
    return _preview_pdf(content)


def preview_paste(text):
    return preview_text(str(text or ""), source_label="Texto colado", format_name="paste")


def annotate_duplicates(result, existing_rows):
    """Acrescenta avisos de duplicidade sem escolher ou alterar nada."""
    existing = {normalized_name(row["name"]): row for row in existing_rows}
    seen = {}
    for index, item in enumerate(result["items"]):
        key = normalized_name(item.get("name"))
        if not key:
            continue
        if key in existing:
            row = existing[key]
            item["duplicate"] = {"kind": "existing", "existing_id": row["id"], "existing_name": row["name"]}
            item["warnings"].append("Já existe uma disciplina com esse nome nesta formação; Ignorar é o padrão.")
        elif key in seen:
            item["duplicate"] = {"kind": "preview", "row": seen[key] + 1}
            item["warnings"].append(f"Nome duplicado da linha {seen[key] + 1}; Ignorar é o padrão.")
        else:
            seen[key] = index
        item["state"] = _state(item)
    result["summary"] = summarize(result["items"])
    return result
